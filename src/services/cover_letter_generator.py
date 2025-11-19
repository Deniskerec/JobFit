from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

load_dotenv()
client = OpenAI()


def generate_cover_letter(resume_text, job_description, user_info):
    """Generate a personalized cover letter using AI"""

    prompt = f"""
You are a professional career coach and expert cover letter writer. Create a compelling, personalized cover letter based on the following information:

CANDIDATE'S RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}

CANDIDATE INFO:
Name: {user_info.get('name', 'John Doe')}
Email: {user_info.get('email', 'email@example.com')}
Phone: {user_info.get('phone', '')}

INSTRUCTIONS:
1. Write a professional cover letter that:
   - Opens with a strong, attention-grabbing introduction
   - Clearly states the position being applied for
   - Highlights 2-3 key achievements from the resume that match the job requirements
   - Shows enthusiasm and cultural fit
   - Explains why the candidate is perfect for this role
   - Includes a call to action
   - Closes professionally

2. Tone: Professional yet personable, confident but not arrogant
3. Length: 3-4 paragraphs, approximately 250-350 words
4. Focus on value proposition: what the candidate can bring to the company

Format the letter with these markers:
- Use [DATE] for today's date placeholder
- Use [COMPANY_NAME] as placeholder for company name
- Use [HIRING_MANAGER] as placeholder for hiring manager's name
- Use [POSITION] for the job title

Create a compelling cover letter that will make the hiring manager want to interview this candidate.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",  # Use gpt-4o-mini which supports chat format
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content


def create_cover_letter_docx(cover_letter_text, user_info, filename="cover_letter.docx"):
    """Create a formatted DOCX file from cover letter text"""

    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add header with contact info
    header_info = doc.add_paragraph()
    header_info.alignment = WD_ALIGN_PARAGRAPH.LEFT

    name_run = header_info.add_run(user_info.get('name', 'Your Name'))
    name_run.bold = True
    name_run.font.size = Pt(14)

    header_info.add_run('\n')

    if user_info.get('email'):
        header_info.add_run(f"{user_info['email']}\n")
    if user_info.get('phone'):
        header_info.add_run(f"{user_info['phone']}\n")
    if user_info.get('linkedin'):
        header_info.add_run(f"{user_info['linkedin']}\n")

    # Add space
    doc.add_paragraph()

    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run(datetime.datetime.now().strftime("%B %d, %Y"))

    # Add space
    doc.add_paragraph()

    # Replace placeholders in cover letter text
    today = datetime.datetime.now().strftime("%B %d, %Y")
    cover_letter_text = cover_letter_text.replace('[DATE]', today)

    # Add cover letter body
    paragraphs = cover_letter_text.split('\n\n')

    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.add_run(para_text.strip())
            para_format = para.paragraph_format
            para_format.space_after = Pt(12)

    # Save document
    output_path = f"/tmp/{filename}"
    doc.save(output_path)

    print(f"Saved cover letter to: {output_path}")  # Debug

    return output_path  # THIS IS CRITICAL - MUST RETURN THE PATH