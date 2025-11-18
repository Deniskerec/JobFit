import json
import re
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()
client = OpenAI()


def build_modification_prompt(cv_text, selected_suggestions):
    """Build prompt to rewrite CV with selected improvements"""
    suggestions_list = "\n".join([f"- {s}" for s in selected_suggestions])

    prompt = f"""
You are a professional CV writer. I will give you a CV and a list of improvements to apply.

ORIGINAL CV:
{cv_text}

IMPROVEMENTS TO APPLY:
{suggestions_list}

Your task:
1. Read the CV carefully
2. Apply each improvement suggestion
3. Keep the person's original experience and facts - DO NOT make up information
4. Maintain a professional tone
5. Keep the same structure (sections like Experience, Education, Skills)

IMPORTANT - Use these formatting markers:
- For headings (like name, section titles): **HEADING: text here**
- For bold text: **text**
- For bullet points: start line with "• "
- For contact info or smaller text: put on separate lines

Return the COMPLETE improved CV text with formatting markers.

Example format:
**HEADING: JOHN SMITH**
john@email.com | +1-555-1234

**HEADING: PROFESSIONAL SUMMARY**
Experienced data engineer with...

**HEADING: EXPERIENCE**
**Data Engineer | Company Name | 2021-Present**
- Built ETL pipelines processing 500GB daily
- Developed Python scripts for automation
- Optimized SQL queries

Make sure to use these markers throughout the CV!
"""
    return prompt


def modify_cv_with_ai(cv_text, selected_suggestions):
    """Send CV to OpenAI for modification"""
    prompt = build_modification_prompt(cv_text, selected_suggestions)

    response = client.responses.create(
        model="gpt-5-nano-2025-08-07",
        input=prompt
    )

    return response.output_text


def create_docx_from_text(text, output_path):
    """Create a DOCX file from formatted text"""
    doc = Document()

    lines = text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Check if it's a heading
        if line.startswith('**HEADING:'):
            heading_text = line.replace('**HEADING:', '').replace('**', '').strip()
            paragraph = doc.add_heading(heading_text, level=1)
            paragraph.runs[0].font.size = Pt(16)
            paragraph.runs[0].font.bold = True
            paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        # Check if it's a bullet point
        elif line.startswith('• '):
            text_content = line[2:]  # Remove bullet
            paragraph = doc.add_paragraph(style='List Bullet')

            # Handle bold text within bullet points
            parts = re.split(r'(\*\*.*?\*\*)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    paragraph.add_run(part)

        # Regular paragraph with potential bold text
        else:
            paragraph = doc.add_paragraph()

            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    paragraph.add_run(part)

    doc.save(output_path)
    return output_path


def modify_cv(cv_text, selected_suggestions, output_filename="improved_resume.docx"):
    """Main function: modify CV with selected suggestions"""
    # Get improved CV text from AI
    improved_text = modify_cv_with_ai(cv_text, selected_suggestions)

    # Create DOCX file with formatting
    output_path = f"/tmp/{output_filename}"
    create_docx_from_text(improved_text, output_path)

    return output_path, improved_text


# Test it
if __name__ == "__main__":
    test_cv = """
    JOHN SMITH
    Data Engineer

    EXPERIENCE
    Data Engineer at TechCorp (2021-Present)
    - Built ETL pipelines
    - Used Python and SQL
    """

    test_suggestions = [
        "Add AWS experience",
        "Quantify achievements with numbers"
    ]

    path, text = modify_cv(test_cv, test_suggestions)
    print(f"Saved to: {path}")
    print(f"\nImproved CV:\n{text}")