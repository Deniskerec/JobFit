import json
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
import re

load_dotenv()
client = OpenAI()


def build_cv_from_info(cv_data):
    """Generate CV text from user-provided information"""

    # Build experience section
    experience_text = ""
    if cv_data.get('experience') and len(cv_data['experience']) > 0:
        exp = cv_data['experience'][0]
        if exp.get('title') and exp.get('title') != 'N/A':
            experience_text = f"""
WORK EXPERIENCE:
{exp['title']} | {exp['company']} | {exp['duration']}
{exp['responsibilities']}
"""

    # Build education section
    education_text = ""
    if cv_data.get('education') and cv_data['education'].get('degree') and cv_data['education']['degree'] != 'N/A':
        edu = cv_data['education']
        education_text = f"""
EDUCATION:
{edu['degree']}
{edu['university']} | {edu['year']}
"""

    prompt = f"""
You are a professional resume writer. Create a polished, professional resume based on the following information:

NAME: {cv_data['name']}
EMAIL: {cv_data['email']}
PHONE: {cv_data['phone']}
LINKEDIN: {cv_data.get('linkedin', '')}

PROFESSIONAL SUMMARY:
{cv_data['summary']}

{experience_text}

{education_text}

SKILLS:
{cv_data['skills']}

Format the resume professionally using these markers:
- For headings: **HEADING: text**
- For bold text: **text**
- For bullet points: start line with "• "

Create a complete, professional resume. If there is no work experience or education provided, focus on skills, summary, and potential. Make it compelling for entry-level positions.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",  # Use gpt-4o-mini which supports chat format
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content


def generate_cv_file(cv_text, filename):
    """Create a DOCX file from CV text with formatting"""
    doc = Document()

    lines = cv_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Heading
        if line.startswith('**HEADING:'):
            heading_text = line.replace('**HEADING:', '').replace('**', '').strip()
            paragraph = doc.add_heading(heading_text, level=1)
            paragraph.runs[0].font.size = Pt(16)
            paragraph.runs[0].font.bold = True

        # Bullet point
        elif line.startswith('• '):
            text_content = line[2:]
            paragraph = doc.add_paragraph(style='List Bullet')

            parts = re.split(r'(\*\*.*?\*\*)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    paragraph.add_run(part)

        # Regular paragraph
        else:
            paragraph = doc.add_paragraph()

            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    paragraph.add_run(part)

    output_path = f"/tmp/{filename}"
    doc.save(output_path)
    return output_path